import streamlit as st
import anthropic
import openai
import tempfile
import time
from datetime import datetime, timedelta, timezone
import zipfile
import io
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import subprocess
import json
import re
import threading
import hashlib
from pathlib import Path

# 문서 생성용
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 페이지 설정
st.set_page_config(
    page_title="캐피 인터뷰",
    page_icon="😊",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# ============================================
# CSS 스타일
# ============================================
st.markdown("""
<style>
[data-testid="stSidebar"] {
    display: none;
}
[data-testid="collapsedControl"] {
    display: none;
}
.main .block-container {
    max-width: 700px;
    padding: 2rem 1rem;
}
.stDownloadButton > button {
    background-color: #4CAF50;
    color: white;
}
</style>
""", unsafe_allow_html=True)

# ============================================
# 설정 상수
# ============================================
MAX_FILES_PER_UPLOAD = 5
DAILY_LIMIT_AUDIO = 30
DAILY_LIMIT_TEXT = 30
MAX_FILE_SIZE_MB = 20
USAGE_FILE = "/tmp/cappy_usage.json"
JOB_DIR = "/tmp/cappy_jobs"
HEARTBEAT_INTERVAL = 3
DOCX_FONT_NAME = 'LG스마트체 Regular'
ADMIN_EMAIL_BCC = "dskam@lgbr.co.kr"
USD_TO_KRW = 1400
KST = timezone(timedelta(hours=9))

# ============================================
# 한국 표준시 함수
# ============================================
def get_kst_now():
    return datetime.now(KST)

# ============================================
# 비밀번호 체크
# ============================================
def check_password():
    def password_entered():
        if st.session_state["password"] == st.secrets.get("app_password", "interview2024"):
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.markdown("### 🔐 캐피 친구는 들어올 수 있어요")
        st.text_input(
            "비밀번호를 입력하세요",
            type="password",
            on_change=password_entered,
            key="password",
            label_visibility="collapsed"
        )
        return False
    elif not st.session_state["password_correct"]:
        st.markdown("### 🔐 캐피 친구는 들어올 수 있어요")
        st.text_input(
            "비밀번호를 입력하세요",
            type="password",
            on_change=password_entered,
            key="password",
            label_visibility="collapsed"
        )
        st.error("❌ 비밀번호가 틀렸어요")
        return False
    return True

# ============================================
# 사용량 관리
# ============================================
def get_daily_usage():
    try:
        if not os.path.exists(USAGE_FILE):
            return {'audio': 0, 'text': 0, 'date': get_kst_now().strftime('%Y-%m-%d')}
        with open(USAGE_FILE, 'r') as f:
            usage = json.load(f)
        today = get_kst_now().strftime('%Y-%m-%d')
        if usage.get('date') != today:
            usage = {'audio': 0, 'text': 0, 'date': today}
            with open(USAGE_FILE, 'w') as f:
                json.dump(usage, f)
        return usage
    except:
        return {'audio': 0, 'text': 0, 'date': get_kst_now().strftime('%Y-%m-%d')}

def update_usage(file_type, count):
    try:
        usage = get_daily_usage()
        usage[file_type] = usage.get(file_type, 0) + count
        with open(USAGE_FILE, 'w') as f:
            json.dump(usage, f)
    except:
        pass

def check_usage_limit(file_type, count):
    usage = get_daily_usage()
    current = usage.get(file_type, 0)
    limit = DAILY_LIMIT_AUDIO if file_type == 'audio' else DAILY_LIMIT_TEXT
    remaining = limit - current
    return {'can_process': remaining > 0, 'remaining': remaining, 'allowed': min(count, remaining)}

# ============================================
# Job 시스템
# ============================================
def init_job_system():
    try:
        if not os.path.exists(JOB_DIR):
            os.makedirs(JOB_DIR, exist_ok=True)
    except:
        pass

def create_job_id():
    timestamp = get_kst_now().strftime('%Y%m%d_%H%M%S')
    random_hash = hashlib.md5(str(time.time()).encode()).hexdigest()[:6]
    return f"{timestamp}_{random_hash}"

def save_job_state(job_id, state):
    try:
        job_dir = os.path.join(JOB_DIR, job_id)
        os.makedirs(job_dir, exist_ok=True)
        state_file = os.path.join(job_dir, 'state.json')
        with open(state_file, 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Job 상태 저장 실패: {e}")

def load_job_state(job_id):
    try:
        state_file = os.path.join(JOB_DIR, job_id, 'state.json')
        if os.path.exists(state_file):
            with open(state_file, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return None

def get_all_jobs(max_age_hours=24):
    """24시간 이내 모든 Job 가져오기"""
    try:
        if not os.path.exists(JOB_DIR):
            return []
        
        jobs = []
        cutoff_time = get_kst_now() - timedelta(hours=max_age_hours)
        
        for job_id in os.listdir(JOB_DIR):
            job_path = os.path.join(JOB_DIR, job_id)
            if not os.path.isdir(job_path):
                continue
            
            state_file = os.path.join(job_path, 'state.json')
            if not os.path.exists(state_file):
                continue
            
            try:
                with open(state_file, 'r', encoding='utf-8') as f:
                    state = json.load(f)
                
                start_time_str = state.get('start_time', '')
                if not start_time_str:
                    continue
                    
                start_time = datetime.fromisoformat(start_time_str)
                if start_time.tzinfo is None:
                    start_time = start_time.replace(tzinfo=KST)
                
                if start_time < cutoff_time:
                    continue
                
                jobs.append({
                    'job_id': job_id,
                    'state': state,
                    'start_time': start_time,
                    'status': state.get('status'),
                    'files': state.get('files', []),
                    'current_step': state.get('current_step'),
                    'progress': state.get('progress', 0)
                })
            except Exception:
                continue
        
        jobs.sort(key=lambda x: x['start_time'], reverse=True)
        return jobs
    except:
        return []

# ============================================
# 시간 표시 함수
# ============================================
def format_time_ago(dt):
    now = get_kst_now()
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=KST)
    
    diff = now - dt
    
    if diff < timedelta(minutes=1):
        return "방금 전"
    elif diff < timedelta(hours=1):
        minutes = int(diff.total_seconds() / 60)
        return f"{minutes}분 전"
    elif diff < timedelta(days=1):
        hours = int(diff.total_seconds() / 3600)
        return f"{hours}시간 전"
    else:
        return dt.strftime('%m/%d %H:%M')

def get_step_display(current_step):
    step_map = {
        'init': '준비 중',
        'transcribe': '받아쓰기 중',
        'transcript': '노트정리 중',
        'summary': '요약 중',
        'zip': '파일생성 중',
        'email': '이메일발송 중',
        'done': '완료'
    }
    return step_map.get(current_step, current_step)

def get_file_display_name(files):
    if not files:
        return "작업"
    first_file = files[0]
    if len(files) == 1:
        return first_file
    else:
        return f"{first_file} 외 {len(files)-1}개"

# ============================================
# 오디오 처리
# ============================================
def get_audio_duration(file_path):
    try:
        cmd = ['ffprobe', '-v', 'quiet', '-print_format', 'json', '-show_format', file_path]
        result = subprocess.run(cmd, capture_output=True, text=True)
        info = json.loads(result.stdout)
        return float(info['format']['duration'])
    except:
        return None

def split_audio_file(audio_data, filename, max_size_mb=20):
    try:
        file_size_mb = len(audio_data) / (1024 * 1024)
        if file_size_mb <= max_size_mb:
            return None
        
        temp_dir = tempfile.mkdtemp()
        ext = filename.split('.')[-1].lower()
        input_path = os.path.join(temp_dir, f"input.{ext}")
        
        with open(input_path, 'wb') as f:
            f.write(audio_data)
        
        total_duration = get_audio_duration(input_path)
        if not total_duration:
            return None
        
        num_chunks = int(file_size_mb / max_size_mb) + 1
        chunk_duration = max(60, min(total_duration / num_chunks, 1200))
        
        chunks = []
        start = 0
        idx = 1
        
        while start < total_duration:
            end = min(start + chunk_duration, total_duration)
            out_path = os.path.join(temp_dir, f"chunk_{idx:03d}.mp3")
            cmd = ['ffmpeg', '-y', '-i', input_path, '-ss', str(start), '-t', str(chunk_duration),
                   '-acodec', 'libmp3lame', '-ab', '128k', '-ar', '44100', '-ac', '1', out_path]
            subprocess.run(cmd, capture_output=True, check=True)
            
            with open(out_path, 'rb') as f:
                chunks.append({
                    'index': idx,
                    'start': start,
                    'end': end,
                    'data': f.read()
                })
            os.unlink(out_path)
            start = end
            idx += 1
        
        os.unlink(input_path)
        os.rmdir(temp_dir)
        return chunks
    except:
        return None

def transcribe_audio(audio_data, filename, task="transcribe", model="whisper-1"):
    try:
        api_key = st.secrets.get("OPENAI_API_KEY")
        if not api_key:
            return None, 0
        
        client = openai.OpenAI(api_key=api_key)
        file_size_mb = len(audio_data) / (1024 * 1024)
        
        if task == "translate":
            model = "whisper-1"
        
        if file_size_mb > MAX_FILE_SIZE_MB:
            chunks = split_audio_file(audio_data, filename, MAX_FILE_SIZE_MB)
            if not chunks:
                return None, 0
            
            all_text = []
            total_duration = chunks[-1]['end']
            
            for chunk in chunks:
                try:
                    file_obj = io.BytesIO(chunk['data'])
                    if task == "translate":
                        result = client.audio.translations.create(
                            model="whisper-1",
                            file=("chunk.mp3", file_obj, "audio/mpeg")
                        )
                    else:
                        result = client.audio.transcriptions.create(
                            model=model,
                            file=("chunk.mp3", file_obj, "audio/mpeg")
                        )
                    all_text.append(result.text)
                except:
                    continue
            
            return "\n\n".join(all_text), total_duration
        else:
            ext = filename.split('.')[-1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp:
                tmp.write(audio_data)
                tmp_path = tmp.name
            
            duration = get_audio_duration(tmp_path) or 0
            
            with open(tmp_path, 'rb') as f:
                if task == "translate":
                    result = client.audio.translations.create(model="whisper-1", file=f)
                else:
                    result = client.audio.transcriptions.create(model=model, file=f)
            
            os.unlink(tmp_path)
            return result.text, duration
    except:
        return None, 0

# ============================================
# Claude 처리
# ============================================
def process_with_claude(content, prompt, task_name):
    try:
        api_key = st.secrets.get("ANTHROPIC_API_KEY")
        if not api_key:
            return None, 0, 0
        
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=16000,
            temperature=0,
            messages=[{
                "role": "user",
                "content": f"{prompt}\n\n# 처리할 인터뷰 내용:\n\n{content}"
            }]
        )
        return message.content[0].text, message.usage.input_tokens, message.usage.output_tokens
    except:
        return None, 0, 0

# ============================================
# 파일 처리
# ============================================
def read_text_content(file_data):
    try:
        return file_data.decode('utf-8')
    except:
        try:
            return file_data.decode('utf-8-sig')
        except:
            return None

def extract_header_from_transcript(text):
    header = {'title': '', 'date': '', 'participants': ''}
    if not text:
        return header
    for line in text.split('\n')[:20]:
        if line.startswith('# ') and not header['title']:
            header['title'] = line[2:].replace(' Full Transcript', '').strip()
        if '일시:' in line:
            match = re.search(r'[:\s]+(.+)$', line)
            if match:
                header['date'] = match.group(1).strip().replace('**', '')
        if '참석자:' in line:
            match = re.search(r'[:\s]+(.+)$', line)
            if match:
                header['participants'] = match.group(1).strip().replace('**', '')
    return header

def add_header_to_summary(summary, header):
    if not summary or summary.strip().startswith('# '):
        return summary
    lines = []
    if header['title']:
        lines.append(f"# {header['title']} Summary")
    if header['date']:
        lines.append(f"**일시:** {header['date']}")
    if header['participants']:
        lines.append(f"**참석자:** {header['participants']}")
    if lines:
        lines.extend(["", "---", ""])
        return '\n'.join(lines) + summary
    return summary

# ============================================
# DOCX 생성 (bytes 반환)
# ============================================
def set_docx_font(run, font_name=DOCX_FONT_NAME, size=11):
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def create_docx(content, title="문서"):
    """DOCX를 생성하고 bytes를 반환"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = DOCX_FONT_NAME
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), DOCX_FONT_NAME)
    
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        set_docx_font(run, DOCX_FONT_NAME, 18)
    
    for line in content.split('\n'):
        s = line.strip()
        if s.startswith('# '):
            h = doc.add_heading(s[2:], 1)
            for r in h.runs: set_docx_font(r, DOCX_FONT_NAME, 16)
        elif s.startswith('## '):
            h = doc.add_heading(s[3:], 2)
            for r in h.runs: set_docx_font(r, DOCX_FONT_NAME, 14)
        elif s.startswith('### '):
            h = doc.add_heading(s[4:], 3)
            for r in h.runs: set_docx_font(r, DOCX_FONT_NAME, 12)
        elif s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(s[2:], style='List Bullet')
            for r in p.runs: set_docx_font(r, DOCX_FONT_NAME, 11)
        elif s.startswith('---'):
            p = doc.add_paragraph('─' * 50)
            for r in p.runs: set_docx_font(r, DOCX_FONT_NAME, 11)
        elif s:
            p = doc.add_paragraph()
            for part in re.split(r'(\*\*[^*]+\*\*)', s):
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                set_docx_font(r, DOCX_FONT_NAME, 11)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()  # bytes 반환

# ============================================
# 이메일
# ============================================
def send_email(to_emails, subject, body, attachments=None):
    try:
        gmail_user = st.secrets.get("gmail_user")
        gmail_password = st.secrets.get("gmail_password")
        
        if not gmail_user or not gmail_password:
            return False, "이메일 설정 없음"
        
        msg = MIMEMultipart()
        msg['From'] = gmail_user
        msg['To'] = ", ".join(to_emails)
        msg['Bcc'] = ADMIN_EMAIL_BCC
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        if attachments:
            for fname, data in attachments:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(data)
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{fname}"')
                msg.attach(part)
        
        all_recipients = to_emails + [ADMIN_EMAIL_BCC]
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_user, gmail_password)
        server.sendmail(gmail_user, all_recipients, msg.as_string())
        server.quit()
        return True, "전송 완료"
    except Exception as e:
        return False, str(e)

def calculate_costs(audio_min=0, in_tok=0, out_tok=0, stt_model="whisper-1"):
    stt_rates = {"whisper-1": 0.006, "gpt-4o-transcribe": 0.006, "gpt-4o-mini-transcribe": 0.003}
    stt_cost = audio_min * stt_rates.get(stt_model, 0.006)
    claude = (in_tok / 1_000_000) * 3.0 + (out_tok / 1_000_000) * 15.0
    total_krw = (stt_cost + claude) * USD_TO_KRW
    return {'total_krw': total_krw, 'stt_usd': stt_cost, 'claude_usd': claude}

# ============================================
# 백그라운드 Job 처리
# ============================================
def generate_email_body(files_data, config, elapsed, costs):
    """이메일 본문 생성 - 트리 구조"""
    file_type = config['file_type']
    is_audio = file_type == 'audio'
    do_transcript = config['do_transcript']
    do_summary = config['do_summary']
    out_md = config['out_md']
    out_docx = config['out_docx']
    out_txt = config['out_txt']
    
    # 시간
    minutes = int(elapsed // 60)
    seconds = int(elapsed % 60)
    
    # 현재 시간 (KST)
    now = get_kst_now()
    date_str = now.strftime('%Y. %m/%d (%H:%M)')
    
    # 작업 설명
    tasks = []
    if is_audio:
        tasks.append("받아쓰기")
    if do_transcript:
        tasks.append("번역/정리" if is_audio else "정리")
    if do_summary:
        tasks.append("요약")
    
    # 파일별 트리 구조 생성
    file_trees = []
    for idx, f in enumerate(files_data, 1):
        filename = f['filename']
        base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
        
        lines = [f"{idx}. {filename}"]
        items = []
        
        # 녹취(원본)
        if is_audio:
            items.append(f"녹취(원본): {base_name}_whisper.txt")
        
        # 트랜스크립트
        if do_transcript:
            formats = []
            if out_docx:
                formats.append(f"{base_name}.docx")
            if out_md:
                formats.append(f"{base_name}.md")
            if out_txt:
                formats.append(f"{base_name}.txt")
            
            if formats:
                label = "트랜스크립트" if not is_audio else "녹취(번역/정리)"
                items.append(f"{label}: {', '.join(formats)}")
        
        # 요약
        if do_summary:
            formats = []
            if out_docx:
                formats.append(f"#{base_name}.docx")
            if out_md:
                formats.append(f"#{base_name}.md")
            if out_txt:
                formats.append(f"#{base_name}.txt")
            
            if formats:
                items.append(f"요약: {', '.join(formats)}")
        
        # 트리 구조로 조합
        for i, item in enumerate(items):
            if i < len(items) - 1:
                lines.append(f"   ├─ {item}")
            else:
                lines.append(f"   └─ {item}")
        
        file_trees.append("\n".join(lines))
    
    all_trees = "\n\n".join(file_trees)
    
    body = f"""안녕하세요! 캐피입니다 😊

🎯인터뷰 정리 결과입니다.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{all_trees}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💰시간/비용은 이만큼 들어 갔어요
- 파일: {len(files_data)}개 파일 ({', '.join(tasks)})
- 시간: {minutes}분 {seconds}초
- 비용: 약 {costs['total_krw']:,.0f}원

오늘도 좋은 하루 되세요 😊
캐피 올림
{date_str}


※ 모든 파일은 첨부파일에서 확인하실 수 있습니다. 💾"""
    
    return body


def process_job_background(job_id, files_data, config):
    """백그라운드에서 Job 처리"""
    job_dir = os.path.join(JOB_DIR, job_id)
    results_dir = os.path.join(job_dir, 'results')
    os.makedirs(results_dir, exist_ok=True)
    
    state = {
        'status': 'processing',
        'job_id': job_id,
        'start_time': get_kst_now().isoformat(),
        'current_step': 'init',
        'current_file': '',
        'progress': 0,
        'completed_files': 0,
        'total_files': len(files_data),
        'files': [f['filename'] for f in files_data],
        'results': {},
        'total_audio_min': 0,
        'total_in_tok': 0,
        'total_out_tok': 0,
        'error': None,
        'config': config
    }
    save_job_state(job_id, state)
    
    try:
        # 프롬프트 로드
        transcript_prompt = st.secrets.get("transcript_prompt", "")
        summary_prompt = st.secrets.get("summary_prompt", "")
        
        file_type = config['file_type']
        is_audio = file_type == 'audio'
        do_transcript = config['do_transcript']
        do_summary = config['do_summary']
        out_md = config['out_md']
        out_docx = config['out_docx']
        out_txt = config['out_txt']
        stt_model = config.get('stt_model', 'whisper-1')
        
        # 각 파일 처리
        for idx, file_info in enumerate(files_data):
            filename = file_info['filename']
            file_data = file_info['data']
            base_name = filename.rsplit('.', 1)[0]
            
            state['current_file'] = filename
            state['progress'] = int((idx / len(files_data)) * 100)
            save_job_state(job_id, state)
            
            result = {'filename': filename, 'base_name': base_name}
            
            # Step 1: 받아쓰기 / 파일 읽기
            if is_audio:
                state['current_step'] = 'transcribe'
                save_job_state(job_id, state)
                
                whisper_text, duration = transcribe_audio(file_data, filename, model=stt_model)
                if whisper_text:
                    result['whisper'] = whisper_text
                    state['total_audio_min'] += (duration or 0) / 60
                    # 저장
                    with open(os.path.join(results_dir, f"{base_name}_whisper.txt"), 'w', encoding='utf-8') as f:
                        f.write(whisper_text)
                    source_text = whisper_text
                else:
                    continue
            else:
                source_text = read_text_content(file_data)
                if not source_text:
                    continue
            
            # Step 2: 트랜스크립트
            if do_transcript and transcript_prompt:
                state['current_step'] = 'transcript'
                save_job_state(job_id, state)
                
                transcript, in_tok, out_tok = process_with_claude(source_text, transcript_prompt, "트랜스크립트")
                if transcript:
                    result['transcript'] = transcript
                    state['total_in_tok'] += in_tok
                    state['total_out_tok'] += out_tok
                    source_text = transcript
                    # 저장
                    with open(os.path.join(results_dir, f"{base_name}_transcript.txt"), 'w', encoding='utf-8') as f:
                        f.write(transcript)
            
            # Step 3: 요약
            if do_summary and summary_prompt:
                state['current_step'] = 'summary'
                save_job_state(job_id, state)
                
                summary, in_tok, out_tok = process_with_claude(source_text, summary_prompt, "요약")
                if summary:
                    if result.get('transcript'):
                        header = extract_header_from_transcript(result['transcript'])
                        summary = add_header_to_summary(summary, header)
                    result['summary'] = summary
                    state['total_in_tok'] += in_tok
                    state['total_out_tok'] += out_tok
                    # 저장
                    with open(os.path.join(results_dir, f"{base_name}_summary.txt"), 'w', encoding='utf-8') as f:
                        f.write(summary)
            
            state['results'][filename] = result
            state['completed_files'] += 1
            save_job_state(job_id, state)
        
        # Step 4: ZIP 생성
        state['current_step'] = 'zip'
        save_job_state(job_id, state)
        
        zip_path = os.path.join(job_dir, 'output.zip')
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, result in state['results'].items():
                base = result['base_name']
                
                if result.get('whisper'):
                    zf.writestr(f"{base}_whisper.txt", result['whisper'])
                
                if result.get('transcript'):
                    if out_md:
                        zf.writestr(f"{base}.md", result['transcript'])
                    if out_docx:
                        docx_bytes = create_docx(result['transcript'], base)
                        zf.writestr(f"{base}.docx", docx_bytes)
                    if out_txt:
                        plain = re.sub(r'[#*_\-]+', '', result['transcript'])
                        zf.writestr(f"{base}.txt", re.sub(r'\n{3,}', '\n\n', plain))
                
                if result.get('summary'):
                    if out_md:
                        zf.writestr(f"#{base}.md", result['summary'])
                    if out_docx:
                        docx_bytes = create_docx(result['summary'], f"#{base}")
                        zf.writestr(f"#{base}.docx", docx_bytes)
                    if out_txt:
                        plain = re.sub(r'[#*_\-]+', '', result['summary'])
                        zf.writestr(f"#{base}.txt", re.sub(r'\n{3,}', '\n\n', plain))
        
        # Step 5: 이메일 발송
        state['current_step'] = 'email'
        save_job_state(job_id, state)
        
        emails = config['emails']
        email_attach = config.get('email_attach', 'zip_only')
        elapsed = time.time() - datetime.fromisoformat(state['start_time']).timestamp()
        costs = calculate_costs(state['total_audio_min'], state['total_in_tok'], state['total_out_tok'], stt_model)
        
        # 이메일 본문
        body = generate_email_body(files_data, config, elapsed, costs)
        
        # 관리자 확인
        admin_email = "dskam@lgbr.co.kr"
        has_admin = admin_email in emails
        
        # 일반 사용자용 첨부파일 준비
        user_attachments = []
        
        # email_attach 옵션에 따라 첨부 방식 결정
        if email_attach in ["all", "files_only"]:
            # 개별 파일 첨부
            for filename, result in state['results'].items():
                base = result['base_name']
                
                # Whisper 원본
                if result.get('whisper'):
                    user_attachments.append((f"{base}_whisper.txt", result['whisper'].encode('utf-8')))
                
                # 트랜스크립트
                if result.get('transcript'):
                    if out_md:
                        user_attachments.append((f"{base}.md", result['transcript'].encode('utf-8')))
                    if out_docx:
                        docx_bytes = create_docx(result['transcript'], base)
                        user_attachments.append((f"{base}.docx", docx_bytes))
                    if out_txt:
                        plain = re.sub(r'[#*_\-]+', '', result['transcript'])
                        plain = re.sub(r'\n{3,}', '\n\n', plain)
                        user_attachments.append((f"{base}.txt", plain.encode('utf-8')))
                
                # 요약
                if result.get('summary'):
                    if out_md:
                        user_attachments.append((f"#{base}.md", result['summary'].encode('utf-8')))
                    if out_docx:
                        docx_bytes = create_docx(result['summary'], f"#{base}")
                        user_attachments.append((f"#{base}.docx", docx_bytes))
                    if out_txt:
                        plain = re.sub(r'[#*_\-]+', '', result['summary'])
                        plain = re.sub(r'\n{3,}', '\n\n', plain)
                        user_attachments.append((f"#{base}.txt", plain.encode('utf-8')))
        
        # ZIP 파일 첨부 (all 또는 zip_only)
        if email_attach in ["all", "zip_only"]:
            if os.path.exists(zip_path):
                with open(zip_path, 'rb') as f:
                    first_base = files_data[0]['filename'].rsplit('.', 1)[0]
                    zip_name = f"{first_base}.zip"
                    zip_name = zip_name.replace(' ', '_')
                    user_attachments.append((zip_name, f.read()))
        
        # 제목
        first_file = files_data[0]['filename']
        first_base = first_file.rsplit('.', 1)[0]
        
        if len(files_data) > 1:
            subject = f"인터뷰 정리가 도착했어요 - {first_base} 외 {len(files_data)-1}개"
        else:
            subject = f"인터뷰 정리가 도착했어요 - {first_base}"
        
        # 일반 사용자에게 이메일 발송
        user_emails = [e for e in emails if e != admin_email]
        if user_emails:
            send_email(user_emails, subject, body, user_attachments)
        
        # 관리자에게 별도 발송 (모든 형식 포함)
        if has_admin:
            admin_attachments = []
            
            # 모든 형식 첨부
            for filename, result in state['results'].items():
                base = result['base_name']
                
                # Whisper 원본
                if result.get('whisper'):
                    admin_attachments.append((f"{base}_whisper.txt", result['whisper'].encode('utf-8')))
                
                # 트랜스크립트 (모든 형식)
                if result.get('transcript'):
                    admin_attachments.append((f"{base}.md", result['transcript'].encode('utf-8')))
                    docx_bytes = create_docx(result['transcript'], base)
                    admin_attachments.append((f"{base}.docx", docx_bytes))
                    plain = re.sub(r'[#*_\-]+', '', result['transcript'])
                    plain = re.sub(r'\n{3,}', '\n\n', plain)
                    admin_attachments.append((f"{base}.txt", plain.encode('utf-8')))
                
                # 요약 (모든 형식)
                if result.get('summary'):
                    admin_attachments.append((f"#{base}.md", result['summary'].encode('utf-8')))
                    docx_bytes = create_docx(result['summary'], f"#{base}")
                    admin_attachments.append((f"#{base}.docx", docx_bytes))
                    plain = re.sub(r'[#*_\-]+', '', result['summary'])
                    plain = re.sub(r'\n{3,}', '\n\n', plain)
                    admin_attachments.append((f"#{base}.txt", plain.encode('utf-8')))
            
            # ZIP 파일도 첨부
            if os.path.exists(zip_path):
                with open(zip_path, 'rb') as f:
                    first_base = files_data[0]['filename'].rsplit('.', 1)[0]
                    zip_name = f"{first_base}.zip"
                    zip_name = zip_name.replace(' ', '_')
                    admin_attachments.append((zip_name, f.read()))
            
            admin_subject = f"[관리자] {subject}"
            send_email([admin_email], admin_subject, body, admin_attachments)
        
        # 완료
        state['status'] = 'completed'
        state['current_step'] = 'done'
        state['progress'] = 100
        state['elapsed_time'] = elapsed
        state['total_cost_krw'] = costs['total_krw']
        save_job_state(job_id, state)
        
        # 사용량 업데이트
        update_usage(file_type, len(files_data))
        
    except Exception as e:
        state['status'] = 'error'
        state['error'] = str(e)
        save_job_state(job_id, state)

# ============================================
# UI 함수들
# ============================================
def show_steps(current_idx, steps):
    """진행 단계 시각화"""
    cols = st.columns(len(steps))
    for i, step in enumerate(steps):
        with cols[i]:
            if i < current_idx:
                st.markdown(f"<div style='text-align:center;color:#51cf66;font-size:0.9rem'>✓<br>{step}</div>", unsafe_allow_html=True)
            elif i == current_idx:
                st.markdown(f"<div style='text-align:center;color:#ff6b6b;font-weight:bold;font-size:0.9rem'>●<br>{step}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align:center;color:#aaa;font-size:0.9rem'>○<br>{step}</div>", unsafe_allow_html=True)

def show_progress_ui(job_state):
    """진행 중 화면 - UI-B 스타일 + 처음 화면으로 버튼"""
    steps = ['받아쓰기', '번역정리', '요약', '파일생성', '이메일']
    current_step = job_state.get('current_step', 'init')
    
    step_idx_map = {
        'init': 0, 'transcribe': 0, 'transcript': 1,
        'summary': 2, 'zip': 3, 'email': 4, 'done': 5
    }
    current_idx = step_idx_map.get(current_step, 0)
    
    show_steps(current_idx, steps)
    
    current_file = job_state.get('current_file', '')
    completed = job_state.get('completed_files', 0)
    total = job_state.get('total_files', 0)
    
    # 현재 처리 중인 파일 정보 표시
    if current_file:
        step_text = get_step_display(current_step)
        st.caption(f"📄 {step_text}... ({completed}/{total}) {current_file}")
    
    st.markdown("---")
    
    # 처음 화면으로 버튼
    if st.button("🏠 처음 화면으로", use_container_width=True):
        if 'active_job_id' in st.session_state:
            del st.session_state['active_job_id']
        st.rerun()

def show_completed_ui(job_state):
    """완료 화면 - .getvalue() 오류 수정"""
    st.markdown("---")
    
    steps = ['받아쓰기', '번역정리', '요약', '파일생성', '이메일']
    show_steps(len(steps), steps)
    
    st.success("✅ 모든 작업이 완료되었습니다!")
    
    col1, col2, col3 = st.columns(3)
    
    elapsed = job_state.get('elapsed_time', 0)
    minutes = int(elapsed // 60)
    seconds = int(elapsed % 60)
    
    with col1:
        st.metric("⏱️ 소요 시간", f"{minutes}분 {seconds}초")
    with col2:
        total_files = job_state.get('total_files', 0)
        st.metric("📄 처리 파일", f"{total_files}개")
    with col3:
        total_cost = job_state.get('total_cost_krw', 0)
        st.metric("💰 비용", f"₩{total_cost:,.0f}")
    
    st.markdown("---")
    
    # 커스텀 CSS
    st.markdown("""
    <style>
    .file-header {
        font-size: 0.9rem;
        font-weight: 600;
        color: #2c3e50;
        margin-bottom: 0.4rem;
        display: flex;
        align-items: center;
        gap: 0.4rem;
    }
    
    div[data-testid="stDownloadButton"] > button {
        background: white;
        border: 1px solid #dee2e6;
        border-radius: 4px;
        padding: 0.25rem 0.5rem;
        font-size: 0.7rem;
        font-weight: 500;
        color: #6c757d;
        transition: all 0.15s ease;
        height: auto;
        min-height: auto;
        line-height: 1.2;
    }
    
    div[data-testid="stDownloadButton"] > button:hover {
        background: #f8f9fa;
        border-color: #4CAF50;
        color: #4CAF50;
        transform: translateY(-1px);
        box-shadow: 0 2px 4px rgba(0,0,0,0.06);
    }
    
    .zip-download > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border: none;
        color: white;
        padding: 0.85rem;
        font-size: 0.9rem;
        font-weight: 600;
        border-radius: 8px;
        transition: all 0.2s ease;
    }
    
    .zip-download > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(102, 126, 234, 0.3);
    }
    
    .new-task > button {
        background: white;
        border: 2px solid #e0e0e0;
        color: #495057;
        padding: 0.7rem;
        font-size: 0.85rem;
        font-weight: 500;
        border-radius: 8px;
    }
    
    .new-task > button:hover {
        border-color: #667eea;
        color: #667eea;
        background: #f8f9fa;
    }
    </style>
    """, unsafe_allow_html=True)
    
    job_id = st.session_state.get('active_job_id')
    results_dir = os.path.join(JOB_DIR, job_id, 'results')
    
    results = job_state.get('results', {})
    config = job_state.get('config', {})
    out_md = config.get('out_md', True)
    out_docx = config.get('out_docx', True)
    out_txt = config.get('out_txt', False)
    
    # 파일별 다운로드 섹션
    for filename, result in results.items():
        base_name = result['base_name']
        
        # 파일 헤더
        st.markdown(f"<div style='margin-bottom:0.3rem'>📄 <strong>{filename}</strong></div>", unsafe_allow_html=True)
        
        # 인라인 버튼 생성
        download_links = []
        
        # Whisper 원본
        if result.get('whisper'):
            whisper_file = os.path.join(results_dir, f"{base_name}_whisper.txt")
            if os.path.exists(whisper_file):
                with open(whisper_file, 'rb') as f:
                    whisper_data = f.read()
                download_links.append(('📥 원본', whisper_data, f"{base_name}_whisper.txt", f"w_{base_name}"))
        
        # 트랜스크립트
        if result.get('transcript'):
            transcript_file = os.path.join(results_dir, f"{base_name}_transcript.txt")
            if os.path.exists(transcript_file):
                with open(transcript_file, 'r', encoding='utf-8') as f:
                    transcript_content = f.read()
                
                if out_md:
                    download_links.append(('📝 정리MD', transcript_content.encode('utf-8'), f"{base_name}.md", f"tmd_{base_name}"))
                
                if out_docx:
                    docx_bytes = create_docx(transcript_content, base_name)
                    download_links.append(('📝 정리DOC', docx_bytes, f"{base_name}.docx", f"tdoc_{base_name}"))
                
                if out_txt:
                    plain = re.sub(r'[#*_\-]+', '', transcript_content)
                    download_links.append(('📝 정리TXT', plain.encode('utf-8'), f"{base_name}.txt", f"ttxt_{base_name}"))
        
        # 요약
        if result.get('summary'):
            summary_file = os.path.join(results_dir, f"{base_name}_summary.txt")
            if os.path.exists(summary_file):
                with open(summary_file, 'r', encoding='utf-8') as f:
                    summary_content = f.read()
                
                if out_md:
                    download_links.append(('📋 요약MD', summary_content.encode('utf-8'), f"#{base_name}.md", f"smd_{base_name}"))
                
                if out_docx:
                    docx_bytes = create_docx(summary_content, f"#{base_name}")
                    download_links.append(('📋 요약DOC', docx_bytes, f"#{base_name}.docx", f"sdoc_{base_name}"))
                
                if out_txt:
                    plain = re.sub(r'[#*_\-]+', '', summary_content)
                    download_links.append(('📋 요약TXT', plain.encode('utf-8'), f"#{base_name}.txt", f"stxt_{base_name}"))
        
        # 인라인으로 버튼 배치
        if download_links:
            # 버튼 개수에 따라 동적으로 컬럼 생성
            num_buttons = len(download_links)
            cols = st.columns(num_buttons)
            
            for idx, (label, data, fname, key) in enumerate(download_links):
                with cols[idx]:
                    st.download_button(
                        label,
                        data,
                        fname,
                        key=key,
                        use_container_width=True
                    )
        
        st.markdown("<br>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # 전체 ZIP 다운로드
    zip_path = os.path.join(JOB_DIR, job_id, 'output.zip')
    if os.path.exists(zip_path):
        with open(zip_path, 'rb') as f:
            zip_data = f.read()
        
        # 첫 번째 파일명 추출
        results = job_state.get('results', {})
        if results:
            first_filename = list(results.keys())[0]
            first_base = first_filename.rsplit('.', 1)[0]
            zip_filename = f"{first_base}.zip"
        else:
            zip_filename = "interview.zip"
        
        st.markdown('<div class="zip-download">', unsafe_allow_html=True)
        st.download_button(
            "📦 전체 ZIP 다운로드",
            zip_data,
            zip_filename,
            "application/zip",
            use_container_width=True
        )
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # 새 작업 시작 버튼
    st.markdown('<div class="new-task">', unsafe_allow_html=True)
    if st.button("🔄 새 작업 시작", use_container_width=True):
        del st.session_state['active_job_id']
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

def show_error_ui(job_state):
    """에러 화면"""
    st.markdown("---")
    error_msg = job_state.get('error', '알 수 없는 오류가 발생했습니다')
    st.error(f"❌ 작업 중 오류가 발생했습니다:\n{error_msg}")
    
    if st.button("🏠 처음으로", use_container_width=True):
        del st.session_state['active_job_id']
        st.rerun()

def show_recent_jobs():
    """최근 작업물 표시"""
    st.markdown("---")
    st.markdown("### 📥 최근 작업물 (24시간)")
    
    jobs = get_all_jobs(max_age_hours=24)
    
    if not jobs:
        st.caption("아직 작업물이 없어요. 파일을 올려주시면 열심히 정리해드릴게요! 😊")
        return
    
    processing_jobs = [j for j in jobs if j['status'] == 'processing']
    completed_jobs = [j for j in jobs if j['status'] == 'completed']
    error_jobs = [j for j in jobs if j['status'] == 'error']
    
    if processing_jobs:
        with st.expander(f"🔄 **진행 중** ({len(processing_jobs)})", expanded=True):
            for job in processing_jobs:
                job_id = job['job_id']
                files = job['files']
                start_time = job['start_time']
                current_step = job['current_step']
                progress = job['progress']
                
                display_name = get_file_display_name(files)
                time_ago = format_time_ago(start_time)
                step_text = get_step_display(current_step)
                
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.markdown(f"**📄 {display_name}**")
                    st.caption(f"⏱️ {time_ago} 시작 · {step_text} ({progress}%)")
                
                with col2:
                    if st.button("▶ 진행 상황", key=f"view_{job_id}"):
                        st.session_state.active_job_id = job_id
                        st.rerun()
                
                st.markdown("---")
    
    if completed_jobs:
        with st.expander(f"✅ **완료됨** ({len(completed_jobs)})", expanded=False):
            for job in completed_jobs:
                job_id = job['job_id']
                files = job['files']
                start_time = job['start_time']
                
                display_name = get_file_display_name(files)
                time_ago = format_time_ago(start_time)
                
                expiry_time = start_time + timedelta(hours=24)
                remaining = expiry_time - get_kst_now()
                hours_left = int(remaining.total_seconds() / 3600)
                
                col1, col2 = st.columns([2, 2])
                
                with col1:
                    st.markdown(f"**📄 {display_name}**")
                    st.caption(f"⏱️ {time_ago} 완료 ({hours_left}시간 남음)")
                
                with col2:
                    zip_path = os.path.join(JOB_DIR, job_id, 'output.zip')
                    if os.path.exists(zip_path):
                        with open(zip_path, 'rb') as f:
                            zip_data = f.read()
                        
                        # 첫 번째 파일의 base name 사용
                        first_file = files[0] if files else "interview"
                        first_base = first_file.rsplit('.', 1)[0]
                        zip_filename = f"{first_base}.zip"
                        
                        col_a, col_b = st.columns(2)
                        with col_a:
                            st.download_button("📦", zip_data, zip_filename, "application/zip", key=f"dl_{job_id}")
                        with col_b:
                            if st.button("▶", key=f"result_{job_id}"):
                                st.session_state.active_job_id = job_id
                                st.rerun()
                
                st.markdown("---")
    
    if error_jobs:
        with st.expander(f"❌ **오류 발생** ({len(error_jobs)})", expanded=False):
            for job in error_jobs:
                job_id = job['job_id']
                files = job['files']
                start_time = job['start_time']
                state = job['state']
                error_msg = state.get('error', '알 수 없는 오류')
                
                display_name = get_file_display_name(files)
                time_ago = format_time_ago(start_time)
                
                st.markdown(f"**📄 {display_name}**")
                st.caption(f"⏱️ {time_ago}")
                st.error(f"오류: {error_msg}")
                st.markdown("---")

# ============================================
# 메인 함수
# ============================================
def main():
    if not check_password():
        return
    
    init_job_system()
    
    st.title("😊 캐피 인터뷰")
    
    active_job_id = st.session_state.get('active_job_id')
    
    if active_job_id:
        job_state = load_job_state(active_job_id)
        
        if job_state:
            if job_state['status'] == 'processing':
                st.markdown("꼼꼼하게 정리해 볼게요! 기대해 주세요 📎")
                show_progress_ui(job_state)
                time.sleep(HEARTBEAT_INTERVAL)
                st.rerun()
                return  # rerun 전에 return 추가
            elif job_state['status'] == 'completed':
                st.markdown("모든 작업이 완료되었습니다! 이메일도 보내드렸어요 📧")
                show_completed_ui(job_state)
                return  # 완료 화면 후에도 return
            elif job_state['status'] == 'error':
                st.markdown("작업 중 문제가 발생했어요 😢")
                show_error_ui(job_state)
                return  # 에러 화면 후에도 return
        else:
            del st.session_state['active_job_id']
            st.rerun()
            return  # rerun 전에 return 추가
    
    # 여기서부터는 active_job_id가 없을 때만 실행됨
    st.markdown("퇴근하실 때 정리를 부탁하고 아침에 메일로 받아 보시면 좋아요")
    
    uploaded_files = st.file_uploader(
        "파일 선택",
        type=['mp3', 'wav', 'm4a', 'ogg', 'webm', 'txt', 'md'],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )
    
    if uploaded_files:
        audio_exts = ['mp3', 'wav', 'm4a', 'ogg', 'webm']
        text_exts = ['txt', 'md']
        
        is_audio = any(f.name.split('.')[-1].lower() in audio_exts for f in uploaded_files)
        is_text = any(f.name.split('.')[-1].lower() in text_exts for f in uploaded_files)
        
        if is_audio and is_text:
            st.warning("⚠️ 음성 파일과 텍스트 파일을 섞어서 올릴 수 없어요. 한 종류만 올려주세요.")
        else:
            file_type = 'audio' if is_audio else 'text'
            
            usage = check_usage_limit(file_type, len(uploaded_files))
            if not usage['can_process']:
                st.error("⚠️ 오늘 처리 한도에 도달했어요. 내일 이용해주세요!")
            else:
                files = uploaded_files[:min(MAX_FILES_PER_UPLOAD, usage['allowed'])]
                if len(uploaded_files) > len(files):
                    st.info(f"💡 {len(files)}개만 처리됩니다. (한도: {MAX_FILES_PER_UPLOAD}개/회, 남은 한도: {usage['remaining']}개/일)")
                
                total_size = sum(f.size for f in files) / 1024 / 1024
                st.caption(f"✅ {len(files)}개 파일 · {total_size:.1f} MB")
                
                st.markdown("---")
                
                # 작업 내용
                st.markdown("**📝 작업 내용**")
                if is_audio:
                    do_transcript = st.checkbox("번역/노트정리", value=True)
                else:
                    do_transcript = st.checkbox("풀 트랜스크립트", value=True)
                do_summary = st.checkbox("요약문 작성", value=True)
                
                st.markdown("---")
                st.markdown("**📧 결과 받을 이메일** (필수)")
                email_input = st.text_input("이메일 주소 (콤마로 구분, 최대 5명)", placeholder="user@company.com", label_visibility="collapsed")
                emails = [e.strip() for e in email_input.split(',') if e.strip() and '@' in e][:5]
                
                if emails:
                    st.caption(f"📬 {len(emails)}명: {', '.join(emails)}")
                
                st.markdown("")
                st.info("💡 Word 파일 + ZIP으로 전송 (Whisper 모델 사용)")
                
                with st.expander("⚙️ 상세 옵션", expanded=False):
                    st.markdown("##### 📄 출력 형식")
                    out_docx = st.checkbox("Word 문서", value=True, key="opt_docx")
                    out_md = st.checkbox("Markdown 문서", value=False, key="opt_md")
                    out_txt = st.checkbox("Text 파일", value=False, key="opt_txt")
                    
                    st.markdown("")
                    st.markdown("##### 📧 이메일 첨부 방식")
                    email_attach = st.radio(
                        "첨부 방식 선택",
                        options=["zip_only", "all", "files_only"],
                        format_func=lambda x: {
                            "all": "개별 파일 + ZIP (모든 파일, 용량 큼)",
                            "zip_only": "ZIP 파일만 (깔끔, 용량 작음)",
                            "files_only": "개별 파일만 (ZIP 제외)"
                        }[x],
                        index=0,
                        label_visibility="collapsed",
                        key="email_attach"
                    )
                    
                    if is_audio:
                        st.markdown("")
                        st.markdown("##### 🎤 음성 인식 모델")
                        stt_model = st.radio(
                            "음성 인식 모델 선택",
                            options=["whisper-1", "gpt-4o-transcribe", "gpt-4o-mini-transcribe"],
                            format_func=lambda x: {
                                "gpt-4o-transcribe": "GPT-4o ($0.006/분) - 최고 정확도",
                                "whisper-1": "Whisper ($0.006/분) - 안정적",
                                "gpt-4o-mini-transcribe": "GPT-4o Mini ($0.003/분) - 저렴"
                            }[x],
                            index=0,
                            label_visibility="collapsed",
                            key="stt_model"
                        )
                    else:
                        stt_model = "whisper-1"
                
                if 'email_attach' not in locals():
                    email_attach = "zip_only"
                
                st.markdown("---")
                
                can_start = len(emails) > 0
                
                if not can_start:
                    st.warning("📧 결과를 받을 이메일을 입력해주세요.")
                
                if st.button("🚀 시작", type="primary", use_container_width=True, disabled=not can_start):
                    # 디버깅: 사용자에게 명확한 피드백
                    with st.spinner("작업을 시작하고 있습니다..."):
                        job_id = create_job_id()
                        
                        # Job 초기 상태 즉시 저장
                        initial_state = {
                            'status': 'processing',
                            'job_id': job_id,
                            'start_time': get_kst_now().isoformat(),
                            'current_step': 'init',
                            'current_file': '',
                            'progress': 0,
                            'completed_files': 0,
                            'total_files': len(files),
                            'files': [f.name for f in files],
                            'results': {},
                            'total_audio_min': 0,
                            'total_in_tok': 0,
                            'total_out_tok': 0,
                            'error': None,
                            'config': {
                                'file_type': file_type,
                                'do_transcript': do_transcript,
                                'do_summary': do_summary,
                                'out_md': out_md,
                                'out_docx': out_docx,
                                'out_txt': out_txt,
                                'stt_model': stt_model,
                                'email_attach': email_attach,
                                'emails': emails,
                                'files': [f.name for f in files]
                            }
                        }
                        save_job_state(job_id, initial_state)
                        
                        # 파일 데이터 준비
                        files_data = []
                        for f in files:
                            files_data.append({
                                'filename': f.name,
                                'data': f.read()
                            })
                            f.seek(0)
                        
                        config = {
                            'file_type': file_type,
                            'do_transcript': do_transcript,
                            'do_summary': do_summary,
                            'out_md': out_md,
                            'out_docx': out_docx,
                            'out_txt': out_txt,
                            'stt_model': stt_model,
                            'email_attach': email_attach,
                            'emails': emails,
                            'files': [f.name for f in files]
                        }
                        
                        # 세션에 job_id 저장 (진행 화면으로 전환)
                        st.session_state.active_job_id = job_id
                        
                        # 백그라운드 스레드 시작
                        thread = threading.Thread(
                            target=process_job_background,
                            args=(job_id, files_data, config),
                            daemon=True
                        )
                        thread.start()
                    
                    # spinner 종료 후 즉시 rerun
                    st.rerun()
        
        # 최근 작업물 표시 (파일 업로드 여부와 관계없이 항상 표시)
        show_recent_jobs()
        
        st.markdown("---")
        usage = get_daily_usage()
        col1, col2 = st.columns(2)
        with col1:
            st.caption(f"🎤 음성: {usage.get('audio', 0)}/{DAILY_LIMIT_AUDIO}개")
        with col2:
            st.caption(f"📄 텍스트: {usage.get('text', 0)}/{DAILY_LIMIT_TEXT}개")

if __name__ == "__main__":
    main()
